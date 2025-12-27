import os
from typing import Any, Dict, List, Tuple
from io import BytesIO

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def _replace_across_runs_preserve_style(
    para: Paragraph,
    old_value: str,
    new_value: str,
    match_case: bool = True,
) -> int:
    if not para.runs or not old_value:
        return 0

    def normalize(s: str) -> str:
        return s if match_case else s.lower()

    run_texts = [r.text or "" for r in para.runs]
    combined = "".join(run_texts)
    combined_norm = normalize(combined)
    old_norm = normalize(old_value)

    if old_norm not in combined_norm:
        return 0

    if match_case:
        occurrences = combined.count(old_value)
        if occurrences == 0:
            return 0
        replaced = combined.replace(old_value, new_value)
    else:
        lower_text = combined_norm
        lower_old = old_norm
        start = 0
        pieces: List[str] = []
        occurrences = 0
        while True:
            idx = lower_text.find(lower_old, start)
            if idx == -1:
                pieces.append(combined[start:])
                break
            pieces.append(combined[start:idx])
            pieces.append(new_value)
            start = idx + len(lower_old)
            occurrences += 1
        if occurrences == 0:
            return 0
        replaced = "".join(pieces)

    remaining = replaced
    for r in para.runs:
        if not remaining:
            r.text = ""
            continue
        orig_len = len(r.text or "")
        r.text = remaining[:orig_len]
        remaining = remaining[orig_len:]

    if remaining:
        para.runs[-1].text = (para.runs[-1].text or "") + remaining

    return occurrences


def _iter_textbox_paragraphs(document: _Document):
    body = document.element.body
    txbx_tag = qn("w:txbxContent")
    p_tag = qn("w:p")

    for txbx in body.iter():
        if txbx.tag != txbx_tag:
            continue
        for p in txbx.iter():
            if p.tag == p_tag:
                yield Paragraph(p, document)


def _apply_single_table_edit(
    table: Table,
    row_idx: int,
    col_idx: int,
    new_value: str,
    old_value: str | None = None,
    match_case: bool = True,
) -> bool:
    """
    Apply a single precise edit to a table cell while preserving
    bullets/numbering and paragraph formatting.

    Behaviour:
    - Do NOT delete paragraphs.
    - If old_value is given: replace occurrences of old_value in the cell's
      paragraphs (one or more) using run-preserving logic.
    - If old_value is None: replace the full text of the first paragraph
      with new_value (using run-preserving logic) and clear text in later
      paragraphs' runs.
    """
    if row_idx >= len(table.rows):
        return False
    row = table.rows[row_idx]
    if col_idx >= len(row.cells):
        return False

    cell = row.cells[col_idx]
    paras = cell.paragraphs
    if not paras:
        # no paragraphs in cell: simple case
        para = cell.add_paragraph(new_value)
        return True

    changed = False

    if old_value is not None:
        # Replace old_value wherever it appears in any paragraph of the cell
        for para in paras:
            count = _replace_across_runs_preserve_style(
                para, old_value, new_value, match_case
            )
            if count > 0:
                changed = True
    else:
        # No old_value: treat it as "set this cell's visible text to new_value"
        # but preserve paragraph/list formatting by rewriting first paragraph
        first_para = paras[0]
        # Use the same helper with the paragraph's full text as old_value
        original_text = "".join(run.text or "" for run in first_para.runs)
        if original_text or new_value:
            # If original_text is empty, treat whole paragraph as target
            target_old = original_text if original_text else original_text + " "
            _replace_across_runs_preserve_style(
                first_para,
                target_old,
                new_value,
                match_case=True,  # here we just overwrite
            )
            # Clear text from other paragraphs but keep them for list formatting
            for para in paras[1:]:
                for run in para.runs:
                    run.text = ""
            changed = True

    return changed


def replace_text_in_document_bytes(
    file_bytes: bytes,
    replacements: List[Tuple[str, str]],
    table_edits: List[Dict[str, Any]] = None,
    match_case: bool = True,
    filename: str | None = None,
) -> Tuple[BytesIO, int]:
    """
    Replace text in one DOCX file (from bytes) in-memory.

    - table_edits: list of dicts like:
        {
          "file": "3.Brouche -springboot-a.docx",  # or suffix
          "table_index": 0,
          "row": 1,
          "col": 1,
          "old_value": "Introduction to spring boot",  # optional
          "new_value": "Introduction to Langchain"
        }
      Only edits whose 'file' matches this document are applied.
      These edits affect ONLY tables.

    - replacements: (old_value, new_value) applied to:
        - main body paragraphs
        - text boxes / shapes
      (tables are skipped; only table_edits affect them)

    - filename: optional, used to filter table_edits by file.

    Returns: (output_stream, total_replacements_count)
    """
    input_stream = BytesIO(file_bytes)
    doc = Document(input_stream)

    total_file_replacements = 0
    basename = os.path.basename(filename) if filename else None

    # 1) Apply precise table edits for this file only
    if table_edits:
        for edit in table_edits:
            try:
                target_file = edit.get("file")
                if target_file and basename:
                    if not (basename == target_file or basename.endswith(target_file)):
                        continue

                table_idx = edit.get("table_index", 0)
                row_idx = edit.get("row", 0)
                col_idx = edit.get("col", 0)
                new_value = edit.get("new_value", "")
                old_value = edit.get("old_value")  # may be None

                if 0 <= table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    applied = _apply_single_table_edit(
                        table=table,
                        row_idx=row_idx,
                        col_idx=col_idx,
                        new_value=new_value,
                        old_value=old_value,
                        match_case=match_case,
                    )
                    if applied:
                        total_file_replacements += 1
            except Exception as e:
                print(f"Error applying table edit for file '{basename}': {e}")
                continue

    # 2) Apply normal text replacements to non-table content
    for old_value, new_value in replacements:
        if not old_value:
            continue

        # main-story paragraphs
        for para in doc.paragraphs:
            total_file_replacements += _replace_across_runs_preserve_style(
                para, old_value, new_value, match_case
            )

        # text boxes / shapes
        for tb_para in _iter_textbox_paragraphs(doc):
            total_file_replacements += _replace_across_runs_preserve_style(
                tb_para, old_value, new_value, match_case
            )

    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    return output_stream, total_file_replacements


if __name__ == "__main__":
    doc_paths = [
        r"C:\Users\Krishna Bhagavan\projects\experiments\docs\1.Request Leatter sB.docx",
        r"C:\Users\Krishna Bhagavan\projects\experiments\docs\3.Brouche -springboot-a.docx",
    ]

    output_paths = [
        "modified_request_letter.docx",
        "modified_brochure.docx",
    ]

    # Normal replacements (won't touch tables)
    multiple_replacements = [
        (
            "Build Web/Enterprise Applications using SpringBoot WITH REST API",
            "Build Application using Generative AI",
        ),
        ("About Guest Lecture", "About Na Lecture"),
        (
            "Artificial Intelligence and Machine Learning",
            "Computer Science Engineering",
        ),
    ]

    # Table edits with file reference (0-based indices)
    table_edits = [
        {
            "file": "1.Request Leatter sB.docx",
            "table_index": 0,
            "row": 2,
            "col": 2,
            "old_value": "0",
            "new_value": "30,000/-",
        },
        {
            "file": "3.Brouche -springboot-a.docx",
            "table_index": 0,
            "row": 1,
            "col": 1,
            # If bullet is paragraph numbering, text is just the label
            "old_value": "Introduction to spring boot",
            "new_value": "Introduction to Langchain",
        },
        {
            "file": "3.Brouche -springboot-a.docx",
            "table_index": 0,
            "row": 2,
            "col": 1,
            "old_value": "Why Microservices ndroid Application development using MIT App Inventor",
            "new_value": "Why Microservices android Application development using MIT App Inventor",
        },
    ]

    try:
        for src, dst in zip(doc_paths, output_paths):
            print(f"\nReading {src}...")
            with open(src, "rb") as f:
                input_bytes = f.read()

            print(f"Processing {os.path.basename(src)} in memory...")
            modified_stream, count = replace_text_in_document_bytes(
                file_bytes=input_bytes,
                replacements=multiple_replacements,
                table_edits=table_edits,
                match_case=True,
                filename=src,
            )

            print(
                f"Total replacements (paragraphs/textboxes + table edits) applied in {os.path.basename(src)}: {count}"
            )
            with open(dst, "wb") as f:
                f.write(modified_stream.getbuffer())

            print(f"Success! Please check '{dst}' to verify changes.")
            print("  - Only table_edits with matching 'file' affected tables")
            print("  - Normal replacements did not change text inside tables")

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")