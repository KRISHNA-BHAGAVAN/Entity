from docx import Document
import io
from typing import List, Dict, Any


def extract_tables_from_docx_bytes(file_bytes: bytes) -> List[Dict[str, Any]]:
    """
    Extract table data from DOCX file bytes.

    For each cell, keep:
      - raw_paragraphs: list of paragraph texts (per bullet/line)
      - display_text: paragraphs joined with '\n' for frontend grid

    The backend can later reconstruct the exact old_value/new_value by
    joining raw_paragraphs with '\n'.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        tables: List[Dict[str, Any]] = []

        for i, table in enumerate(doc.tables):
            table_preview: List[List[str]] = []
            table_paragraphs: List[List[List[str]]] = []  # [row][col][para_text]

            for row in table.rows:
                row_preview: List[str] = []
                row_paragraphs: List[List[str]] = []

                for cell in row.cells:
                    paras = [p.text or "" for p in cell.paragraphs]

                    # Normalise whitespace, but keep paragraph boundaries
                    cleaned_paras = [p.strip() for p in paras if p.strip() != ""]

                    if not cleaned_paras:
                        display_text = ""
                    else:
                        # Show as multi-line text in grid; frontend can split on '\n'
                        display_text = "\n".join(cleaned_paras)

                    row_preview.append(display_text)
                    row_paragraphs.append(cleaned_paras)

                table_preview.append(row_preview)
                table_paragraphs.append(row_paragraphs)

            if table_preview:
                tables.append({
                    "index": i,
                    "rows": len(table_preview),
                    "columns": len(table_preview[0]) if table_preview else 0,
                    "preview": table_preview,
                    # raw paragraphs per cell so you can reconstruct full text
                    "paragraphs": table_paragraphs,
                })

        return tables
    except Exception as e:
        print(f"Error extracting tables: {e}")
        return []
    
if __name__ == "__main__":
    import json
    import os
    
    # Update this path to your test DOCX file
    test_docx_path = r"C:\Users\Krishna Bhagavan\projects\experiments\docs\3.Brouche -springboot-a.docx"
    
    print("=== Testing extract_tables_from_docx_bytes ===")
    
    # Test file-based extraction
    if os.path.exists(test_docx_path):
        print(f"Testing file: {os.path.basename(test_docx_path)}")
        with open(test_docx_path, "rb") as f:
            file_bytes = f.read()
        
        tables = extract_tables_from_docx_bytes(file_bytes)
        print(f"Found {len(tables)} tables")
        
        for i, table in enumerate(tables):
            print(f"\nTable {table['index']}: {table['rows']}x{table['columns']}")
            print("Preview:")
            for row_idx, row in enumerate(table['preview']):  # First 3 rows
                print(f"  Row {row_idx}: {row}")
            # if len(table['preview']) > 3:
            #     print(f"  ... and {len(table['preview'])-3} more rows")
    else:
        print(f"File not found: {test_docx_path}")
        print("Please update test_docx_path to point to a DOCX file with tables")
    
    print("\nTest complete!")
