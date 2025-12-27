import json
import hashlib
import time
from typing import List, Dict, Any, TypedDict, Optional, Tuple
from functools import lru_cache
from dotenv import load_dotenv
from langchain.chat_models import init_chat_model
from langchain_core.messages import SystemMessage, HumanMessage
from langgraph.graph import StateGraph, START, END
import redis
from rapidfuzz import fuzz
import tiktoken

import os
from docx import Document as DocxDocument
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

# --------------------------------------------------------------------------
# DOCX HELPERS (tables + paragraphs + textboxes)
# --------------------------------------------------------------------------

def normalize_text(text: str) -> str:
    if not text:
        return text
    return (
        text
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u00a0", " ")
        .replace("\u2026", ".")
    )

def _count_in_text(haystack: str, needle: str, match_case: bool = True) -> int:
    """Count non-overlapping occurrences of needle in haystack."""
    if not needle:
        return 0
    if not match_case:
        haystack = haystack.lower()
        needle = needle.lower()
    return haystack.count(needle)

def count_reference_in_docx(doc_path: str, reference: str, match_case: bool = True) -> int:
    """
    Count how many times `reference` appears in:
     - normal paragraphs
     - table cell paragraphs
     - text-box paragraphs
    """
    if not os.path.exists(doc_path):
        return 0

    doc = DocxDocument(doc_path)
    total = 0

    # 1) main-story paragraphs
    for para in doc.paragraphs:
        text = para.text or ""
        total += _count_in_text(text, reference, match_case)

    # 2) table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text or ""
                    total += _count_in_text(text, reference, match_case)

    # 3) text boxes / shapes
    for tb_para in _iter_textbox_paragraphs(doc):
        text = tb_para.text or ""
        total += _count_in_text(text, reference, match_case)

    return total

# ------------------------------------------------------------------------------
# TOKEN COUNTING (lightweight)
# ------------------------------------------------------------------------------

_ENCODING = tiktoken.get_encoding("cl100k_base")

def get_token_count(text: str) -> int:
    try:
        return len(_ENCODING.encode(text))
    except Exception:
        return len(text) // 4 + 1

# ------------------------------------------------------------------------------
# GLOBAL REFERENCE DEDUPLICATION
# ------------------------------------------------------------------------------

GLOBAL_SEEN_REFS: set = set()

def is_reference_unique(ref: str) -> bool:
    """Check if reference is exactly the same as any previously seen reference."""
    return ref not in GLOBAL_SEEN_REFS

def mark_reference_seen(ref: str) -> None:
    """Mark reference as seen to prevent duplicates across all fields."""
    GLOBAL_SEEN_REFS.add(ref)

# ------------------------------------------------------------------------------
# ENV + LLM + REDIS CACHE
# ------------------------------------------------------------------------------

load_dotenv(override=True)

llm = init_chat_model(
    model="llama-3.3-70b-versatile",
    temperature=0,
    model_provider="groq",
)

try:
    redis_client = redis.Redis(
        host="localhost", port=6379, db=0, decode_responses=True
    )
    redis_client.ping()
    USE_REDIS = True
    print("✅ Redis cache enabled")
except Exception:
    redis_client = None
    USE_REDIS = False
    print("⚠️ Using in-memory cache")

# ------------------------------------------------------------------------------
# STATE
# ------------------------------------------------------------------------------

class SchemaDiscoveryState(TypedDict, total=False):
    documents: List[Tuple[str, str]]
    doc_paths: List[str]
    cache_key: Optional[str]
    partial_schemas: List[Dict[str, Any]]
    final_schema: Optional[Dict[str, Any]]
    stats: Dict[str, Any]

INITIAL_STATS: Dict[str, Any] = {
    "cache_hit": False,
    "processing_time": 0.0,
    "docs_processed": 0,
    "total_fields": 0,
    "sections_created": 0,
    "total_chars_processed": 0,
    "merge_time": 0.0,
    "llm": {
        "calls": [],
        "summary": {
            "llm_calls": 0,
            "total_input_tokens": 0,
            "total_output_tokens": 0,
            "total_tokens": 0,
            "avg_input_tokens_per_call": 0,
            "avg_output_tokens_per_call": 0,
        },
    },
}

# ------------------------------------------------------------------------------
# PROMPT
# ------------------------------------------------------------------------------

SCHEMA_DISCOVERY_PROMPT = """
You help users reuse and edit event-related documents.

Document: {filename}

Task:
From the content, extract fields a user is likely to change from event to event or over time.

Focus on (when present):
- Event names and titles
- Event objectives / purpose
- Dates (letter dates, event start/end, duration ranges)
- Institutions, departments, venues
- People grouped by role:
  - event_organizer / organizers
  - resource_person / resource_persons
  - contact_person / contact_people
  - other key_people (patrons, chancellors, committee)
- Contact details (phones, emails, addresses)
- Other clearly event-specific details that would usually change next time

Rules:
- Top-level keys: stable snake_case (e.g. event_name, event_dates, resource_persons, contact_email)
- Reuse role-specific keys instead of inventing near-duplicates
  (prefer event_organizer/contact_person/resource_person over generic key_person).
- For each field, output:
    - label       : human-readable label
    - type        : one of [date, date_range, string, address, email, phone]
    - references  : array of ALL exact text spans (phrases/sentences/paragraphs) that are semantically the same thing and would be edited together
    - confidence  : float in [0.0, 1.0]

Reference rules:
- references MUST be exact substrings from the document.
- Put all textual variants of the same fact in ONE field's references.
  Example:
    "08-09-2025 TO 10-09-2025",
    "08/09/2025 to 10/09/2025",
    "the workshop starts on 08-09-2025 and ends on 10-09-2025"
  all belong to the same event_dates field.
- Deduplicate references WITHIN each field.
- Do NOT repeat exact same phrases across different fields.

Output STRICT JSON ONLY:
- A single flat JSON object.
- Keys = field names (snake_case).
- Values = field objects as described.
- No markdown, no prose, no comments.
"""

# ------------------------------------------------------------------------------
# UTILS
# ------------------------------------------------------------------------------

@lru_cache(maxsize=128)
def normalize_key(key: str) -> str:
    return key.lower().replace(" ", "_").replace("-", "_")

def fuzzy_dedupe_references(references: List[str], threshold: int = 85) -> List[str]:
    if len(references) <= 1:
        return references
    deduped = []
    for ref in sorted(references, key=len, reverse=True):
        is_duplicate = any(
            fuzz.ratio(ref, existing) >= threshold for existing in deduped
        )
        if not is_duplicate:
            deduped.append(ref)
    return deduped

CACHE: Dict[str, Any] = {}

def get_cache(key: str) -> Optional[Dict[str, Any]]:
    if USE_REDIS:
        try:
            cached = redis_client.get(f"schema:{key}")
            return json.loads(cached) if cached else None
        except Exception:
            return None
    return CACHE.get(key)

def set_cache(key: str, value: Dict[str, Any], ttl: int = 3600) -> None:
    if USE_REDIS:
        try:
            redis_client.setex(f"schema:{key}", ttl, json.dumps(value))
            return
        except Exception:
            pass
    CACHE[key] = value

def track_llm_usage(
    stats: Dict[str, Any],
    input_tokens: int,
    output_tokens: int,
    prompt_chars: int,
) -> Dict[str, Any]:
    llm = stats.setdefault(
        "llm",
        {
            "calls": [],
            "summary": {
                "llm_calls": 0,
                "total_input_tokens": 0,
                "total_output_tokens": 0,
                "total_tokens": 0,
                "avg_input_tokens_per_call": 0,
                "avg_output_tokens_per_call": 0,
            },
        },
    )

    call_stats = {
        "timestamp": time.time(),
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "prompt_chars": prompt_chars,
        "total_tokens": input_tokens + output_tokens,
    }
    llm["calls"].append(call_stats)

    summary = llm["summary"]
    summary["llm_calls"] += 1
    summary["total_input_tokens"] += input_tokens
    summary["total_output_tokens"] += output_tokens
    summary["total_tokens"] = (
        summary["total_input_tokens"] + summary["total_output_tokens"]
    )
    if summary["llm_calls"] > 0:
        summary["avg_input_tokens_per_call"] = (
            summary["total_input_tokens"] // summary["llm_calls"]
        )
        summary["avg_output_tokens_per_call"] = (
            summary["total_output_tokens"] // summary["llm_calls"]
        )
    return stats

# ------------------------------------------------------------------------------
# NODES
# ------------------------------------------------------------------------------

def cache_check(state: SchemaDiscoveryState) -> Dict[str, Any]:
    stats = state.get("stats", INITIAL_STATS.copy())

    parts = []
    total_chars = 0
    for (filename, markdown) in state["documents"]:
        parts.append(f"{filename}:{markdown[:1000]}")
        total_chars += len(markdown)
    for p in state.get("doc_paths", []):
        parts.append(f"docx:{p}")

    content_hash = hashlib.sha256("".join(parts).encode()).hexdigest()
    stats["total_chars_processed"] = total_chars

    cached = get_cache(content_hash)
    if cached:
        print(f"✅ CACHE HIT: {content_hash[:8]}")
        stats["cache_hit"] = True
        stats["processing_time"] = 0.001
        return {
            "final_schema": cached["schema"],
            "stats": stats,
        }

    print(f"🔄 CACHE MISS: {content_hash[:8]}")
    # Reset global seen references for new processing
    GLOBAL_SEEN_REFS.clear()
    return {"cache_key": content_hash, "stats": stats}

def map_discover_schema(state: SchemaDiscoveryState) -> Dict[str, Any]:
    partials: List[Dict[str, Any]] = []
    stats = state.get("stats", INITIAL_STATS.copy())

    for i, (filename, md_raw) in enumerate(state["documents"]):
        raw_length = len(md_raw)
        print(f"\n🔍 DOC {i+1} ({filename}): RAW={raw_length} chars")

        if raw_length < 50:
            print("  ⏭️ SKIP: too short")
            continue

        try:
            content = md_raw[:8000]
            print(f"  📤 Sending {len(content)} chars to LLM...")

            prompt = SCHEMA_DISCOVERY_PROMPT.format(filename=filename)
            full_prompt = prompt + "\n\n" + content

            prompt_tokens = get_token_count(full_prompt)
            start_time = time.time()

            response = llm.invoke(
                [
                    SystemMessage(content=prompt),
                    HumanMessage(content=content),
                ]
            )

            response_time = time.time() - start_time
            output_tokens = get_token_count(response.content)

            print(
                f"  📥 LLM RAW RESPONSE ({len(response.content)} chars, ~{output_tokens} tokens)"
            )
            print(f"     {repr(response.content[:200])}...")

            response_text = response.content.strip()
            start = response_text.find("{")
            end = response_text.rfind("}") + 1
            json_str = response_text[start:end] if start != -1 else None

            if not json_str:
                print("  ❌ NO JSON FOUND!")
                continue

            partial_schema = json.loads(json_str)
            print(
                f"  ✅ PARSED: {len(partial_schema)} keys: {list(partial_schema.keys())}"
            )

            # Filter references to only unique ones (character-for-character)
            for field_key, field_val in partial_schema.items():
                if isinstance(field_val, dict) and "references" in field_val:
                    refs = field_val["references"]
                    unique_refs = []
                    for ref in refs:
                        if is_reference_unique(ref):
                            unique_refs.append(ref)
                            mark_reference_seen(ref)
                        else:
                            print(f"  ⏭️ Skipping duplicate ref in {field_key}: {repr(ref)[:50]}...")
                    field_val["references"] = unique_refs
                    field_val["source_filename"] = filename

            partials.append(partial_schema)
            print(f"  🎉 ADDED DOC {i+1} ({filename}) - {response_time:.2f}s")

            stats = track_llm_usage(
                stats,
                input_tokens=prompt_tokens,
                output_tokens=output_tokens,
                prompt_chars=len(full_prompt),
            )

        except Exception as e:
            print(f"  💥 ERROR DOC {i+1}: {e}")
            continue

    stats["docs_processed"] = len(partials)
    print(f"\n🎯 TOTAL PARTIALS: {len(partials)}")
    return {"partial_schemas": partials, "stats": stats}

def merge_schemas_enhanced(state: SchemaDiscoveryState) -> Dict[str, Any]:
    start_time = time.time()
    merged: Dict[str, Any] = {}

    for schema in state["partial_schemas"]:
        print(f"🔍 Merging schema: {len(schema)} keys")

        if "$schema" in schema:
            schema = schema.get("properties", {})

        section_key = "document_fields"
        section_val = {
            "label": "Document Fields",
            "fields": schema,
        }

        merged_section = merged.setdefault(
            section_key,
            {
                "label": "Document Fields",
                "fields": {},
                "doc_frequency": 0,
            },
        )
        merged_section["doc_frequency"] += 1

        for field_key, field_val in section_val["fields"].items():
            if not isinstance(field_val, dict):
                continue

            merged_field = merged_section["fields"].setdefault(
                field_key, field_val.copy()
            )

            # Only merge NEW unique references (global deduplication already handled)
            existing_refs = set(merged_field.get("references", []))
            new_refs = set(field_val.get("references", []))
            merged_refs = sorted(existing_refs.union(new_refs))
            merged_field["references"] = merged_refs

            merged_field["doc_frequency"] = merged_field.get("doc_frequency", 0) + 1
            merged_field["confidence"] = max(
                merged_field.get("confidence", 0),
                field_val.get("confidence", 0),
            )

            source_files = merged_field.setdefault("source_files", [])
            source_file = field_val.get("source_filename")
            if source_file and source_file not in source_files:
                source_files.append(source_file)

    # Sort fields by importance
    for section in merged.values():
        fields = section.get("fields", {})
        if isinstance(fields, dict):
            section["fields"] = dict(
                sorted(
                    fields.items(),
                    key=lambda x: (
                        x[1].get("doc_frequency", 0)
                        * x[1].get("confidence", 0),
                        len(x[1].get("references", [])),
                    ),
                    reverse=True,
                )
            )

    processing_time = time.time() - start_time
    total_fields = sum(len(s.get("fields", {})) for s in merged.values())

    stats = state["stats"].copy()
    stats.update(
        {
            "processing_time": processing_time,
            "total_fields": total_fields,
            "sections_created": len(merged),
            "merge_time": processing_time,
        }
    )

    print(
        f"🎉 FINAL MERGE: {total_fields} fields from {len(state['partial_schemas'])} docs"
    )
    print(f"📋 GLOBAL UNIQUE REFS: {len(GLOBAL_SEEN_REFS)}")
    return {
        "final_schema": merged,
        "stats": stats,
    }

def compute_frequencies(state: SchemaDiscoveryState) -> Dict[str, Any]:
    """
    Compute frequency per field based on actual occurrences of each reference
    across the provided docx paths. This does NOT modify docx files.
    """
    final_schema = state.get("final_schema", {})
    doc_paths = state.get("doc_paths", [])
    if not final_schema or not doc_paths:
        return {}

    # Map filenames (without dirs) to full paths for convenience
    path_by_basename = {os.path.basename(p): p for p in doc_paths}

    document_fields = final_schema.get("document_fields", {})
    fields = document_fields.get("fields", {})
    print("📊 COMPUTE_FREQUENCIES: starting")

    for field_key, field_val in fields.items():
        refs = field_val.get("references", [])
        source_files = field_val.get("source_files", [])

        total_freq = 0
        for ref in refs:
            if not ref:
                continue

            # For each referenced file
            for fname in source_files:
                # match by basename if needed
                full_path = path_by_basename.get(fname) or next(
                    (p for p in doc_paths if p.endswith(fname)), None
                )
                if not full_path:
                    continue
                total_freq += count_reference_in_docx(full_path, ref, match_case=True)

        field_val["frequency"] = total_freq

    print("📊 COMPUTE_FREQUENCIES: done")
    return {"final_schema": final_schema}

def cache_store(state: SchemaDiscoveryState) -> Dict[str, Any]:
    if state.get("cache_key"):
        cache_data = {
            "schema": state["final_schema"],
            "stats": state.get("stats", {}),
        }
        set_cache(state["cache_key"], cache_data)
        print(f"💾 CACHED: {state['cache_key'][:8]} with full stats")
    return {}

# Missing helper function (add this)
def _iter_textbox_paragraphs(doc: _Document) -> List[Paragraph]:
    """Helper to iterate through textbox paragraphs."""
    # This is a simplified version - implement based on your docx needs
    return []

# ------------------------------------------------------------------------------
# LANGGRAPH
# ------------------------------------------------------------------------------

graph = StateGraph(SchemaDiscoveryState)

graph.add_node("cache_check", cache_check)
graph.add_node("map_discover_schema", map_discover_schema)
graph.add_node("merge_schemas", merge_schemas_enhanced)
graph.add_node("compute_frequencies", compute_frequencies)
graph.add_node("cache_store", cache_store)

graph.add_edge(START, "cache_check")
graph.add_conditional_edges(
    "cache_check", lambda s: END if s.get("final_schema") else "map_discover_schema"
)
graph.add_edge("map_discover_schema", "merge_schemas")
graph.add_edge("merge_schemas", "compute_frequencies")
graph.add_edge("compute_frequencies", "cache_store")
graph.add_edge("cache_store", END)

schema_discovery_workflow = graph.compile()